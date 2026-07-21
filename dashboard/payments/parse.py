#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
snEco · payments — розбір архіву рахунків на оплату → структуровані реквізити + призначення платежу.

Джерело: ZIP/тека з рахунками (pdf / скан-pdf / docx / xls(x) / зображення).
Вихід : рядки pmt_invoices (отримувач, ЄДРПОУ, IBAN, №, дата, сума, ПДВ, призначення, прапорці)
        → JSON та/або POST /api/dashboard/ingest (entity=pmt_invoices) з X-Sync-Key.

Платник (наш): ТОВ «ПРАЙМ СНЕК», ЄДРПОУ 40271201.

Формат призначення (вимога НБУ + внутрішня):
  "Оплата за <послуги> згідно рахунку №<номер> від <дата>, у т.ч. ПДВ 20% - <сума ПДВ> грн"
  Неплатник ПДВ → "..., без ПДВ".
  Якщо постачальник диктує призначення / обов'язковий код (напр. ОККО 7400002034) — беремо його + дописуємо ПДВ.

Використання:
  python parse.py --zip рахунки.zip --out out.json --dry-run
  python parse.py --dir ./inbox --ingest            # потрібні env WORKER_URL + SYNC_API_KEY
Опційно: ANTHROPIC_API_KEY → уточнення складних/низькоточних рахунків через Claude.
"""
import re, os, io, sys, json, zipfile, hashlib, argparse, subprocess, tempfile

PAYER_EDRPOU = "40271201"
PAYER_HINTS  = ["прайм снек", "прайм снэк", PAYER_EDRPOU]

# ─────────────────────────── text extraction (multi-format) ───────────────────────────
def _run(cmd, **kw):
    try: return subprocess.run(cmd, capture_output=True, timeout=120, **kw)
    except Exception: return None

def _pdf_text(path):
    r = _run(["pdftotext", "-layout", "-q", path, "-"])
    return r.stdout.decode("utf-8", "ignore") if r and r.returncode == 0 else ""

def _ocr(path_or_imgdir, is_pdf=True):
    txt = ""
    with tempfile.TemporaryDirectory() as td:
        if is_pdf:
            _run(["pdftoppm", "-r", "300", "-png", path_or_imgdir, os.path.join(td, "p")])
            imgs = sorted(f for f in os.listdir(td) if f.endswith(".png"))
            imgs = [os.path.join(td, f) for f in imgs]
        else:
            imgs = [path_or_imgdir]
        for img in imgs:
            r = _run(["tesseract", img, "stdout", "-l", "ukr+rus"])
            if r and r.returncode == 0:
                txt += r.stdout.decode("utf-8", "ignore") + "\n"
    return txt

def _docx_text(path):
    try:
        import docx
        return "\n".join(p.text for p in docx.Document(path).paragraphs) + "\n" + \
               "\n".join(c.text for t in docx.Document(path).tables for r in t.rows for c in r.cells)
    except Exception:
        r = _run(["docx2txt", path]) or _run(["antiword", path])
        return r.stdout.decode("utf-8", "ignore") if r and r.returncode == 0 else ""

def _xlsx_text(path):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        out = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                out.append(" ".join("" if c is None else str(c) for c in row))
        return "\n".join(out)
    except Exception:
        return ""

def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    ocr = False
    if ext == ".pdf":
        t = _pdf_text(path)
        if len(re.sub(r"\s", "", t)) < 40:
            t = _ocr(path, is_pdf=True); ocr = True
    elif ext in (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"):
        t = _ocr(path, is_pdf=False); ocr = True
    elif ext in (".docx",):
        t = _docx_text(path)
    elif ext in (".xlsx", ".xlsm", ".xls"):
        t = _xlsx_text(path)
    else:
        t = ""
    return t, ocr

# ─────────────────────────── field extraction (rule-based) ───────────────────────────
def norm_amount(tok):
    if tok is None: return None
    t = tok.replace("\xa0", "").replace(" ", "").strip().rstrip(".")
    if "," in t and "." in t:
        t = t.replace(".", "").replace(",", ".") if t.rfind(",") > t.rfind(".") else t.replace(",", "")
    elif "," in t:
        t = t.replace(",", ".")
    try: return round(float(t), 2)
    except: return None

NUM = r"\d[\d \xa0]*(?:[.,]\d{2,4})?"
def last_num(s):
    m = re.findall(NUM, s); return norm_amount(m[-1]) if m else None
def all_line_nums(text, labels):
    v = []
    for lab in labels:
        for m in re.finditer(lab + r"[^\n]*", text, re.I):
            x = last_num(m.group(0));  v.append(x) if x is not None else None
    return v

def iban_valid(iban):
    iban = (iban or "").upper()
    if not re.fullmatch(r"UA\d{27}", iban): return False
    r = iban[4:] + iban[:4]
    return int("".join(str(int(c, 36)) for c in r)) % 97 == 1

def find_ibans(text):
    out, seen = [], set()
    for m in re.finditer(r"UA[\d \xa0]{27,45}", text):
        d = re.sub(r"[ \xa0]", "", m.group(0))[:29]
        if len(d) == 29 and d not in seen:
            seen.add(d); out.append((d, m.start()))
    return out

def near(text, pos, hints, w=110):
    seg = text[max(0, pos - w):pos + 10].lower()
    return any(h in seg for h in hints)

def clean_name(s):
    s = re.split(r"\s{2,}", s.strip())[0].strip(" :\t")
    for stop in ["IBAN", "П/р", "п/р", "р/р", "Р/р", "р. №", "код", "Код", "Адреса", "адреса", "тел", "Тел", ","]:
        i = s.find(stop)
        if i > 3: s = s[:i]
    return s.strip(" :,-\t")

def extract_name(text, labels):
    for lab in labels:
        m = re.search(lab + r"\s*:?\s*(.+)", text)
        if m:
            nm = clean_name(m.group(1))
            if len(nm) >= 3 and PAYER_EDRPOU not in nm and "прайм снек" not in nm.lower():
                return nm
    return None

def extract_edrpou(text):
    codes = [m.group(1) for m in re.finditer(r"(?:ЄДРПОУ|ЗКПО|Код)\D{0,8}(\d{8})", text)]
    return next((c for c in codes if c != PAYER_EDRPOU), None), (PAYER_EDRPOU in codes)

def extract_total_vat(text):
    total = None
    for lab in [r"Усього\s*з\s*ПДВ", r"Всього\s*з\s*ПДВ", r"Разом\s*з\s*ПДВ",
                r"Усього\s*до\s*сплати", r"Всього\s*до\s*сплати", r"Разом\s*до\s*сплати", r"До\s*сплати"]:
        m = re.search(lab + r"[^\n]*", text, re.I)
        if m and last_num(m.group(0)): total = last_num(m.group(0)); break
    if total is None:
        m = re.search(r"на\s+сум[уи]\s+(" + NUM + r")", text, re.I)
        if m: total = norm_amount(m.group(1))
    base = None
    m = re.search(r"(?:Всього|Усього|Разом)\s*без\s*ПДВ[^\n]*", text, re.I)
    if m: base = last_num(m.group(0))
    vat = None
    for lab in [r"тому\s*числ\S*\s*ПДВ", r"Сума\s*ПДВ", r"т\.?\s*ч\.?\s*ПДВ", r"^\s*ПДВ\s*[:\-]"]:
        for m in re.finditer(lab + r"[^\n]*", text, re.I | re.M):
            x = last_num(m.group(0))
            if x is not None and x > 0: vat = x; break
        if vat is not None: break
    if total is None and base is not None and vat is not None:
        total = round(base + vat, 2)
    if total is None:
        c = all_line_nums(text, [r"Всього", r"Усього", r"Разом"])
        if c: total = max(c)
    rate = None
    rm = re.search(r"ПДВ\s*(\d{1,2})\s*%", text)
    if rm: rate = int(rm.group(1))
    no_vat = bool(re.search(r"не\s*є\s*платником\s*ПДВ|ПДВ\s*не\s*оподатков|єдин\S*\s*податк|на\s+сум[уи][^\n]*без\s*ПДВ|0\s*%\s*ПДВ", text, re.I))
    derived = False
    if total:
        if vat and 0 < vat < total * 0.5: vat_final = vat; rate = rate or 20
        elif no_vat: vat_final = 0; rate = 0
        else: vat_final = round(total / 6, 2); rate = rate or 20; derived = True
    else:
        vat_final = vat
    return total, vat_final, rate, derived

def extract_invnum(text, fname):
    tnum = None
    m = re.search(r"[Рр]ахун[а-яіїє\-]*(?:-фактура)?\s*(?:№|N)\s*([^\s]+(?:\s*-\s*\d+)?)\s*в[іi]д", text)
    if m: tnum = re.sub(r"\s+", "", m.group(1)).strip()
    fid = re.search(r"(\d{7,})", fname); num = tnum
    if fid and (num is None or fid.group(1) not in (num or "")):
        num = f"{num}-{fid.group(1)}" if num else fid.group(1)
    if num is None:
        fn = re.search(r"№?\s*([A-ZА-ЯЇІ]{0,4}[\-/]?\d{2,}[\-/\d]*)", fname); num = fn.group(1) if fn else None
    multi = bool(re.search(r"\d{2,}\s*,\s*\d{2,}", fname)) or len(set(re.findall(r"[Рр]ахунок\s*(?:№|N)\s*(\d{3,})", text))) > 1
    return num, multi

def extract_purpose_code(text):
    m = re.search(r"(\d{8,12})\s+Оплата\s+зг", text)
    if m: return m.group(1)
    m = re.search(r"призначенн\S*\s*платежу[^\d]{0,40}(\d{8,12})", text, re.I)
    return m.group(1) if m else None

def extract_specified_purpose(text):
    m = re.search(r"[Пп]ризначенн\S*\s*платежу\s*:?\s*\n?\s*([^\n]{6,180})", text)
    if m:
        line = re.sub(r"\s+", " ", m.group(1)).strip().strip('"').lstrip("|•*- ").strip()
        if re.search(r"[Оо]плат|\d{6,}", line) and "вказувати" not in line.lower():
            return line
    return None

SERVICE_MAP = [("маркетинг", "маркетингові послуги"), ("логіст", "логістичні послуги"),
    ("стимулюванн", "послуги зі стимулювання збуту"), ("просуванн", "послуги з просування товару"),
    ("ретро", "ретро-бонус"), ("бонус", "ретро-бонус"), ("оренд", "оренда"),
    ("транспорт", "транспортні послуги"), ("реклам", "рекламні послуги"), ("мерчанд", "мерчандайзинг")]
def service_summary(text):
    low = text.lower()
    for k, v in SERVICE_MAP:
        if k in low: return v
    m = re.search(r"1\s+([А-Яа-яЇїІіЄєҐґЁё][^\n]{6,60})", text)
    if m:
        cand = re.sub(r"\s+", " ", m.group(1)).strip()
        if not re.match(r"^\W*\d", cand) and not re.match(r"(?i)(січн|лют|берез|квіт|трав|черв|лип|серп|верес|жовт|листоп|груд)", cand):
            return cand[:55]
    return "товари/послуги"

def build_purpose(num, date, service, vat, rate, code, spec):
    if spec:
        base = spec.rstrip(" .")
        if code and code not in base: base = f"{code} {base}"
    elif code:
        base = f"{code} Оплата згідно рахунку №{num} від {date}"
    else:
        base = f"Оплата за {service} згідно рахунку №{num} від {date}"
    if not re.search(r"т\.?\s*ч\.?\s*ПДВ|тому числ\S*\s*ПДВ|без ПДВ", base, re.I):
        base += (f", у т.ч. ПДВ {rate or 20}% - {('%.2f' % vat).replace('.', ',')} грн") if (vat and vat > 0) else ", без ПДВ"
    return (base[:157] + "...") if len(base) > 160 else base

def norm_date(text, fname):
    months = {"січня":"01","лютого":"02","березня":"03","квітня":"04","травня":"05","червня":"06",
              "липня":"07","серпня":"08","вересня":"09","жовтня":"10","листопада":"11","грудня":"12"}
    mon = "|".join(months)
    m = re.search(r"(\d{2})\.(\d{2})\.(\d{4})", fname)
    if m: return f"{m.group(1)}.{m.group(2)}.{m.group(3)}"
    m = re.search(r"(\d{1,2})\s+(" + mon + r")\s+(\d{4})", fname)
    if m: return f"{int(m.group(1)):02d}.{months[m.group(2)]}.{m.group(3)}"
    m = re.search(r"(?:рахун\S*|фактур\S*)[^\n]{0,60}?в[іi]д\s*(\d{2})\.(\d{2})\.(\d{4})", text, re.I)
    if m: return f"{m.group(1)}.{m.group(2)}.{m.group(3)}"
    m = re.search(r"в[іi]д\s*(\d{1,2})\s+(" + mon + r")\s+(\d{4})", text, re.I)
    if m: return f"{int(m.group(1)):02d}.{months[m.group(2)]}.{m.group(3)}"
    m = re.search(r"(\d{2})\.(\d{2})\.(\d{4})", text)
    return f"{m.group(1)}.{m.group(2)}.{m.group(3)}" if m else None

def severity(flags):
    if not flags: return "ok"
    if any(h in flags for h in ["IBAN не знайдено", "кілька рахунків", "не пройшов", "сума не знайдена"]): return "red"
    if any(a in flags for a in ["кілька IBAN", "скан", "OCR", "обчислено", "платник"]): return "amber"
    return "ok"

def parse_one(fname, text, ocr):
    flags = []
    name = extract_name(text, ["Постачальник", "Продавець", "Виконавець"]) or extract_name(text, ["Отримувач"])
    rec_edrpou, payer_seen = extract_edrpou(text)
    ibans = find_ibans(text)
    valid = [(i, p) for (i, p) in ibans if iban_valid(i)]
    invalid = [i for (i, p) in ibans if not iban_valid(i)]
    rec_iban = None
    cand = [(i, p) for (i, p) in valid if not near(text, p, PAYER_HINTS)]
    if len(cand) == 1: rec_iban = cand[0][0]
    elif len(cand) > 1: rec_iban = cand[0][0]; flags.append(f"кілька IBAN({len(cand)})")
    elif len(valid) == 1: rec_iban = valid[0][0]
    if not rec_iban and invalid: flags.append("IBAN не пройшов контр.суму")
    if not rec_iban and not ibans: flags.append("IBAN не знайдено")
    num, multi = extract_invnum(text, fname)
    if multi: flags.append("кілька рахунків у файлі—розбити")
    date = norm_date(text, fname)
    total, vat, rate, derived = extract_total_vat(text)
    code = extract_purpose_code(text)
    if code: flags.append(f"обов.код {code}")
    if ocr: flags.append("скан/OCR—звірити")
    if not (payer_seen or "прайм снек" in text.lower() or "прайм снэк" in text.lower()): flags.append("платник?")
    if total is None: flags.append("сума не знайдена")
    if derived: flags.append("ПДВ обчислено 20%")
    spec = extract_specified_purpose(text)
    purpose = build_purpose(num or "?", date or "?", service_summary(text), vat, rate, code, spec)
    fl = "; ".join(flags)
    return {
        "file": fname, "recipient": name, "edrpou": rec_edrpou, "iban": rec_iban,
        "iban_valid": 1 if (rec_iban and iban_valid(rec_iban)) else 0,
        "invoice_no": num, "invoice_date": date,
        "amount_kop": int(round(total * 100)) if total is not None else None,
        "vat_kop": int(round(vat * 100)) if vat is not None else None,
        "vat_rate": rate, "currency": "UAH", "purpose": purpose,
        "severity": severity(fl), "flags": fl,
    }

# ─────────────────────────── optional Claude refinement ───────────────────────────
def claude_refine(row, text):
    """Уточнення складних/низькоточних рахунків через Claude (якщо є ANTHROPIC_API_KEY).
    Активується лише для рядків severity!=ok, щоб не витрачати токени."""
    key = os.environ.get("ANTHROPIC_API_KEY")
    if not key or row["severity"] == "ok":
        return row
    import urllib.request
    schema_hint = ('Витягни з тексту рахунку JSON: recipient, edrpou(8цифр), iban(UA+27цифр, отримувача-продавця, '
        'НЕ платника ЄДРПОУ 40271201), invoice_no, invoice_date(DD.MM.YYYY), amount(з ПДВ, число), '
        'vat(сума ПДВ, число), vat_rate(20/7/0). Якщо у файлі кілька рахунків — поверни масив invoices[].')
    body = {"model": "claude-3-5-sonnet-latest", "max_tokens": 1024,
            "messages": [{"role": "user", "content": f"{schema_hint}\n\nТЕКСТ РАХУНКУ:\n{text[:6000]}\n\nПоверни ЛИШЕ JSON."}]}
    try:
        req = urllib.request.Request("https://api.anthropic.com/v1/messages",
            data=json.dumps(body).encode(), headers={
                "content-type": "application/json", "x-api-key": key, "anthropic-version": "2023-06-01"})
        resp = json.loads(urllib.request.urlopen(req, timeout=60).read())
        txt = resp["content"][0]["text"]
        data = json.loads(re.search(r"\{.*\}|\[.*\]", txt, re.S).group(0))
        inv = data["invoices"][0] if isinstance(data, dict) and "invoices" in data else data
        if isinstance(inv, list): inv = inv[0]
        for k_src, k_dst, conv in [("recipient","recipient",str),("edrpou","edrpou",str),("iban","iban",str),
                ("invoice_no","invoice_no",str),("invoice_date","invoice_date",str)]:
            if inv.get(k_src): row[k_dst] = conv(inv[k_src])
        if inv.get("amount"): row["amount_kop"] = int(round(float(inv["amount"]) * 100))
        if inv.get("vat") is not None: row["vat_kop"] = int(round(float(inv["vat"]) * 100))
        if inv.get("vat_rate") is not None: row["vat_rate"] = int(inv["vat_rate"])
        row["iban_valid"] = 1 if iban_valid(row.get("iban")) else 0
        row["flags"] = (row["flags"] + "; уточнено Claude").strip("; ")
        row["severity"] = severity(row["flags"].replace("уточнено Claude", ""))
    except Exception as e:
        row["flags"] = (row["flags"] + f"; Claude:err").strip("; ")
    return row

# ─────────────────────────── driver ───────────────────────────
def fix_zip_name(n):
    for enc in ("cp866", "cp1251", "utf-8"):
        try: return n.encode("cp437").decode(enc)
        except Exception: continue
    return n

def iter_files(src):
    if src.lower().endswith(".zip"):
        z = zipfile.ZipFile(src)
        for n in z.namelist():
            if n.endswith("/"): continue
            fx = fix_zip_name(n)
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(fx)[1]) as tf:
                tf.write(z.read(n)); tmp = tf.name
            yield os.path.basename(fx), tmp
    else:
        for root, _, fns in os.walk(src):
            for fn in fns:
                if fn.startswith("."): continue
                yield fn, os.path.join(root, fn)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--zip"); ap.add_argument("--dir"); ap.add_argument("--out")
    ap.add_argument("--batch-id", default="local"); ap.add_argument("--uploaded-by", default="")
    ap.add_argument("--ingest", action="store_true"); ap.add_argument("--dry-run", action="store_true")
    ap.add_argument("--claude", action="store_true", help="уточнювати складні рядки через Claude")
    a = ap.parse_args()
    src = a.zip or a.dir
    if not src: sys.exit("need --zip or --dir")

    rows = []
    for i, (fname, path) in enumerate(iter_files(src)):
        text, ocr = extract_text(path)
        row = parse_one(fname, text, ocr)
        if a.claude: row = claude_refine(row, text)
        row["id"] = f"{a.batch_id}:{i:03d}:" + hashlib.md5(fname.encode()).hexdigest()[:8]
        row["batch_id"] = a.batch_id
        rows.append(row)

    total_kop = sum(r["amount_kop"] or 0 for r in rows)
    by_sev = {s: sum(1 for r in rows if r["severity"] == s) for s in ("ok", "amber", "red")}
    print(f"parsed {len(rows)} files | {by_sev} | total {total_kop/100:,.2f} UAH", file=sys.stderr)

    if a.out:
        json.dump({"batch_id": a.batch_id, "rows": rows}, open(a.out, "w"), ensure_ascii=False, indent=1)
        print("wrote", a.out, file=sys.stderr)

    if a.ingest and not a.dry_run:
        import urllib.request
        worker = os.environ["WORKER_URL"].rstrip("/"); key = os.environ["SYNC_API_KEY"]
        def post(entity, payload):
            req = urllib.request.Request(f"{worker}/api/dashboard/ingest",
                data=json.dumps(payload).encode(),
                headers={"content-type": "application/json", "X-Sync-Key": key})
            return json.loads(urllib.request.urlopen(req, timeout=60).read())
        print(post("pmt_batches", {"entity": "pmt_batches", "rows": [{
            "id": a.batch_id, "filename": os.path.basename(src), "uploaded_by": a.uploaded_by,
            "file_count": len(rows), "total_kop": total_kop, "status": "parsed"}]}), file=sys.stderr)
        print(post("pmt_invoices", {"entity": "pmt_invoices", "rows": rows}), file=sys.stderr)

    return rows

if __name__ == "__main__":
    main()
